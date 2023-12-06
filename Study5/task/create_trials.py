import pandas as pd
from docx import Document

dict_second_actoin={'bell.png_train.png':2,'bell.png_tophat.png':8,
'tree.png_north.png':2,'tree.png_snorkel.png':8,
'watch.png_houses.png':2,'watch.png_compass.png':8,
'fox.png_hammer.png':2,'fox.png_apple.png':8}

df=pd.read_csv('Study2_learningtrials_r3.csv')

action2=[]
unicorn_action=2
for row in range(len(df)):
	s2_image=df['s2_image'][row]
	s3_image=df['s3_image'][row]

	if s2_image=='unicorn.png':
		if unicorn_action==2:
			unicorn_action=8
			action2.append(unicorn_action)
		elif unicorn_action==8:
			unicorn_action=2
			action2.append(unicorn_action)
	else:
		action2.append(dict_second_actoin[s2_image+'_'+s3_image])

df['action2']=action2

# df.to_csv('new_trials.csv')

# Calculate and print the percentage of times each 's2_image' state occurred
total_s2_images = len(df)
s2_image_counts = df['s2_image'].value_counts().to_dict()
print(s2_image_counts)
print("Percentage of times each 's2_image' state occurred:")
for s2_image, count in s2_image_counts.items():
    percentage = (count / total_s2_images) * 100
    print(f"{s2_image}: {percentage:.2f}%")


# Initialize a nested dictionary to count 's3_image' occurrences for each 's2_image' and 'action2'
s2_action_to_s3_counts = {}

# Populate the nested dictionary with counts
for index, row in df.iterrows():
    s2_image = row['s2_image']
    s3_image = row['s3_image']
    action2 = row['action2']
    
    state_action_pair = (s2_image, action2)
    
    if state_action_pair not in s2_action_to_s3_counts:
        s2_action_to_s3_counts[state_action_pair] = {}
    
    if s3_image not in s2_action_to_s3_counts[state_action_pair]:
        s2_action_to_s3_counts[state_action_pair][s3_image] = 0
    
    s2_action_to_s3_counts[state_action_pair][s3_image] += 1

# Calculate and write the percentages to a Word document
doc = Document()
doc.add_heading("Percentage of transitions from each ('s2_image', 'action2') to each possible 's3_image':", level=1)

for (s2_image, action2), s3_counts in s2_action_to_s3_counts.items():
    total_transitions = sum(s3_counts.values())
    doc.add_heading(f"({s2_image}, action2='{action2}') (total transitions: {total_transitions}):", level=2)
    for s3_image, count in s3_counts.items():
        percentage = (count / total_transitions) * 100
        doc.add_paragraph(f"{s3_image}: {percentage:.2f}%")

# Save the document
doc.save("transitions.docx")