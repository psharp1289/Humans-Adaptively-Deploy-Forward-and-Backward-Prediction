import pandas as pd
from docx import Document

dict_second_action={'snorkel.png':0,
'north.png':0,
'train.png':0,
'tophat.png':0,
'houses.png':0,
'compass.png':0,
'apple.png':0,
'hammer.png':0,
'window.png':0}


dict_starting_states={}

df=pd.read_csv('just_s1.csv')

medium=['snorkel.png','train.png','north.png','tophat.png']
high=['compass.png']
low=['apple.png','hammer.png']

unicorn_action='a'
fox_action='a'

action2=[]

		
for row in range(len(df)):
	s1_image=df['s1_image'][row]
	dict_second_action[s1_image]+=1
	if s1_image in medium:
		if dict_second_action[s1_image]>39:
			df['s2_image'][row]='watch.png'
			if dict_second_action[s1_image]>59:
				df['s3_image'][row]='bell.png'
				action2.append('a')
			else:
				df['s3_image'][row]='trident.png'
				action2.append('k')
		else:
			df['s2_image'][row]='tree.png'
			if dict_second_action[s1_image]>38:
				df['s3_image'][row]='bell.png'
				action2.append('k')
			else:
				df['s3_image'][row]='planet.png'
				action2.append('a')

	if s1_image in high:
		if dict_second_action[s1_image]<134:
			df['s2_image'][row]='watch.png'
			if dict_second_action[s1_image]>130:
				df['s3_image'][row]='bell.png'
				action2.append('a')
			else:
				df['s3_image'][row]='trident.png'
				action2.append('k')
		else:
			df['s2_image'][row]='tree.png'
			if dict_second_action[s1_image]>200:
				df['s3_image'][row]='bell.png'
				action2.append('k')
			else:
				df['s3_image'][row]='planet.png'
				action2.append('a')
	if s1_image in low:
			df['s2_image'][row]='fox.png'
			df['s3_image'][row]='trident.png'
			action2.append('k')
			
	
df['action2']=action2
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)
df = df.sample(frac=1).reset_index(drop=True)

# df.to_csv('easier_probability_trials.csv',index=False)


# Calculate and print the percentage of times each 's2_image' state occurred
total_s2_images = len(df)
s2_image_counts = df['s2_image'].value_counts().to_dict()

print("Percentage of times each 's2_image' state occurred:")
for s2_image, count in s2_image_counts.items():
    percentage = (count / total_s2_images) * 100
    print(f"{s2_image}: {percentage:.2f}%")


# Initialize a nested dictionary to count 's3_image' occurrences for each 's2_image' and 'action2'
s2_action_to_s3_counts = {}

# Populate the nested dictionary with counts
for index, row in df.iterrows():
    s2_image = row['s2_image']
    s3_image = row['s3_image']
    action2 = row['action2']
    
    state_action_pair = (s2_image, action2)
    
    if state_action_pair not in s2_action_to_s3_counts:
        s2_action_to_s3_counts[state_action_pair] = {}
    
    if s3_image not in s2_action_to_s3_counts[state_action_pair]:
        s2_action_to_s3_counts[state_action_pair][s3_image] = 0
    
    s2_action_to_s3_counts[state_action_pair][s3_image] += 1

# Calculate and write the percentages to a Word document
doc = Document()
doc.add_heading("Percentage of transitions from each ('s2_image', 'action2') to each possible 's3_image':", level=1)

for (s2_image, action2), s3_counts in s2_action_to_s3_counts.items():
    total_transitions = sum(s3_counts.values())
    doc.add_heading(f"({s2_image}, action2='{action2}') (total transitions: {total_transitions}):", level=2)
    for s3_image, count in s3_counts.items():
        percentage = (count / total_transitions) * 100
        doc.add_paragraph(f"{s3_image}: {percentage:.2f}%")

# Save the document
doc.save("output.docx")